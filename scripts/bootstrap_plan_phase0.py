"""Validate the upgrade workspace and materialize the DOCX plan into Phase 0 reports."""
from __future__ import annotations

import json
import subprocess
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
PLAN = Path(r"C:\Users\phong\Downloads\ICTAI_FIX\KE_HOACH_NANG_CAP_PAPER_ICTAI2026_CHO_CODEX_GIU_DATASET_V1.docx")
REFERENCE_DIR = PLAN.parent

TASKS = [
    ("P0-01", "Validate workspace and read-only plan path", "reports/path_validation.*", "PASS"),
    ("P0-02", "Extract plan and establish task/decision registries", "reports/plan_extracted.*, reports/master_task_board.*", "PASS"),
    ("P0-03", "Inventory repository and freeze Dataset V1", "data/legacy/state_panel_v1/V1_IMMUTABLE_MANIFEST.json", "PENDING"),
    ("P1-01", "Verify frozen Dataset V1 before new experiments", "legacy-v1 verification report", "BLOCKED_BY_P0-03"),
    ("P2-01", "Implement synthetic fidelity-gate benchmark", "artifacts/experiments/synthetic-gate-benchmark", "BLOCKED_BY_P1-01"),
    ("P2-02", "Implement county-level Dataset V2 pipeline", "data/v2_county", "BLOCKED_BY_P1-01"),
    ("P2-03", "Run feature/model ladders and promotion gates", "reports/experiments", "BLOCKED_BY_P2-02"),
    ("P3-01", "External-domain validation or blocker report", "reports/experiments/external-domain.*", "BLOCKED_BY_P1-01"),
    ("P4-01", "Score applied/method/hybrid routes and revise paper", "reports/final_route_scorecard.*", "BLOCKED_BY_P2-01,P2-02,P3-01"),
]


def command(*args: str) -> str:
    return subprocess.check_output(args, cwd=ROOT, text=True, encoding="utf-8", errors="replace").strip()


def main() -> None:
    reports = ROOT / "reports"
    reports.mkdir(exist_ok=True)
    if not ROOT.is_dir() or not (ROOT / ".git").exists():
        raise SystemExit("BLOCKED_WORKSPACE_PATH")
    if not PLAN.is_file():
        raise SystemExit("BLOCKED_PLAN_DOCUMENT")
    document = Document(PLAN)
    paragraphs = [
        {"style": paragraph.style.name, "text": paragraph.text.strip()}
        for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    tables = [
        [[cell.text.strip() for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]
    git = {"branch": command("git", "branch", "--show-current"), "commit": command("git", "rev-parse", "HEAD"), "status": command("git", "status", "--short")}
    path_payload = {
        "status": "PASS", "workspace_root": str(ROOT), "plan_docx": str(PLAN), "resolved_workspace_root": str(ROOT.resolve()), "resolved_plan_docx": str(PLAN.resolve()), "current_working_directory": str(Path.cwd().resolve()), "repository_root": command("git", "rev-parse", "--show-toplevel"), "git": git, "plan_exists_and_readable": True, "reference_location_read_only": str(REFERENCE_DIR), "writes_outside_workspace": False,
    }
    (reports / "path_validation.json").write_text(json.dumps(path_payload, indent=2) + "\n", encoding="utf-8")
    (reports / "path_validation.md").write_text("# Path Validation\n\n" + "\n".join(f"- {key.replace('_', ' ')}: `{value}`" for key, value in path_payload.items()) + "\n", encoding="utf-8")
    extraction = {"extracted_utc": datetime.now(timezone.utc).isoformat(), "plan_docx": str(PLAN), "paragraphs": paragraphs, "tables": tables}
    (reports / "plan_extracted.json").write_text(json.dumps(extraction, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    lines = ["# Extracted Upgrade Plan", "", f"- Source: `{PLAN}`", "", "## Paragraphs", ""]
    lines.extend(f"- **{item['style']}**: {item['text']}" for item in paragraphs)
    for index, table in enumerate(tables, 1):
        lines.extend(["", f"## Table {index}", ""])
        lines.extend(" | ".join(row) for row in table)
    (reports / "plan_extracted.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    board = [{"task_id": task_id, "objective": objective, "output": output, "dependency": "see status", "branch": "main until experiment branches are created", "command": "not run in Phase 0" if status != "PASS" else "python scripts/bootstrap_plan_phase0.py", "expected_runtime": "TBD after freeze", "result": "Phase 0 evidence written" if status == "PASS" else "not started", "status": status, "next_action": "Freeze and verify Dataset V1" if task_id == "P0-03" else "Wait for dependencies"} for task_id, objective, output, status in TASKS]
    (reports / "master_task_board.json").write_text(json.dumps(board, indent=2) + "\n", encoding="utf-8")
    table_lines = ["# Master Task Board", "", "| ID | Objective | Output | Status | Next action |", "|---|---|---|---|---|"]
    table_lines.extend(f"| {row['task_id']} | {row['objective']} | `{row['output']}` | {row['status']} | {row['next_action']} |" for row in board)
    (reports / "master_task_board.md").write_text("\n".join(table_lines) + "\n", encoding="utf-8")
    decisions = "# Decision Registry\n\n- D-001: The DOCX plan is read-only and remains outside the repository. Status: ACCEPTED.\n- D-002: No V2, synthetic, external-domain, or model experiments may run until Dataset V1 is frozen and verified. Status: ENFORCED.\n- D-003: Dataset V1 remains the immutable state-level baseline; candidate work requires separate versioned artifacts. Status: ACCEPTED.\n"
    (reports / "decision_registry.md").write_text(decisions, encoding="utf-8")
    print("Phase 0 path validation and plan extraction PASS")


if __name__ == "__main__":
    main()
